from docx import Document
from collections import Counter
from pathlib import Path

path = Path('data/documents/PIS-2116_Location Based Service_A-V0.0.2.3.docx')
doc = Document(path)
styles = Counter()
for para in doc.paragraphs:
    name = para.style.name if para.style else '<None>'
    styles[name] += 1

print('Total paragraphs:', len(doc.paragraphs))
for name, count in styles.most_common(30):
    print(f'{name}: {count}')
