from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from langchain_community.document_loaders import Docx2txtLoader, PyPDFLoader
from langchain_core.documents import Document

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    print("⚠️  pdfplumber 未安装，PDF 将使用降级加载器。安装: pip install pdfplumber")

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("⚠️  pandas 未安装，无法加载 Excel 文件。安装: pip install pandas openpyxl")


# ==================== 文档类型检测配置 ====================
DOC_TYPE_MAP = {
    # nav_spec - 导航功能规范
    'PIS-2116_Location Based Service': 'nav_spec',
    'PIS-2114_Data Security': 'nav_spec',
    'NDLB Map Service': 'nav_spec',

    # api_spec - API接口规范
    'AI灵图FSA接口文档': 'api_spec',
    'Launcher接口说明': 'api_spec',
    'SDK说明文档': 'api_spec',

    # process_spec - 流程规范
    'Bug fix': 'process_spec',
    '代码提交': 'process_spec',
    '分支策略': 'process_spec',
    '高德地图推包': 'process_spec',
    '报AI灵图bug': 'process_spec',

    # metrics_spec - 指标规范
    'AutoHelp': 'metrics_spec',
    '性能指标': 'metrics_spec',

    # doc_generic - 通用文档
    'AI灵图使命': 'doc_generic',
}


def detect_doc_type(path: Path) -> str:
    """
    根据文件名检测文档类型

    Args:
        path: 文档路径

    Returns:
        文档类型: nav_spec | api_spec | process_spec | metrics_spec | doc_generic
    """
    filename = path.stem  # 不含扩展名的文件名

    # 遍历映射字典，匹配关键词
    for keyword, doc_type in DOC_TYPE_MAP.items():
        if keyword in filename:
            return doc_type

    # 默认类型
    return 'doc_generic'


@dataclass
class _SectionNode:
    """层级节点，用于构建文档的树状结构"""
    title: str  # 完整标题文本（包含节号，如 "3.2.1 LBS系统介绍"）
    level: int  # 层级深度 (1, 2, 3...)
    section_number: Optional[str] = None  # 节号 (如 "3.2.1")
    section_title: Optional[str] = None  # 纯标题文本 (如 "LBS系统介绍")
    contents: List[str] = field(default_factory=list)
    children: List["_SectionNode"] = field(default_factory=list)


def _parse_section_number(text: str) -> tuple[Optional[str], str]:
    """
    从标题文本中解析节号和纯标题

    支持的格式：
    - "3.2.1 标题" → ("3.2.1", "标题")
    - "3.2.1.1 标题" → ("3.2.1.1", "标题")
    - "第3章 标题" → ("第3章", "标题")
    - "(1) 标题" → ("(1)", "标题")
    - "标题" → (None, "标题")

    Args:
        text: 标题文本

    Returns:
        (节号, 纯标题文本)
    """
    text = text.strip()

    # 模式1: 数字编号 (3.2.1 标题)
    match = re.match(r'^([\d.]+)\s+(.+)$', text)
    if match:
        return match.group(1), match.group(2)

    # 模式2: 中文章节 (第3章 标题)
    match = re.match(r'^(第[一二三四五六七八九十\d]+[章节])\s*(.*)$', text)
    if match:
        return match.group(1), match.group(2) or match.group(1)

    # 模式3: 括号编号 ((1) 标题)
    match = re.match(r'^(\([一二三四五1-9]+\))\s*(.+)$', text)
    if match:
        return match.group(1), match.group(2)

    # 无节号
    return None, text


def _parse_heading_level(style_name: Optional[str]) -> Optional[int]:
    if not style_name:
        return None
    name = style_name.lower()
    match = re.search(r'(\d+)', name)
    if match and ('heading' in name or '标题' in style_name):
        return int(match.group(1))
    return None


def _build_loader(path: Path):
    suffix = path.suffix.lower()
    if suffix == '.pdf':
        return PyPDFLoader(str(path))
    if suffix == '.docx':
        return Docx2txtLoader(str(path))
    raise ValueError(f'Unsupported file type: {path.suffix}')


def _load_docx_with_headings(path: Path) -> List[Document]:
    from docx import Document as WordDocument

    word_doc = WordDocument(str(path))
    root = _SectionNode(title='root', level=0)
    stack = [root]

    # --- 新增：辅助方法，把表格转成结构化文本 ---
    def extract_table(table):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # 跳过空行
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            return "\n".join(rows)
        return None

    # --- 遍历整个 document.body，包含 paragraphs 和 tables ---
    for block in word_doc.element.body:
        if block.tag.endswith('p'):  # 段落
            para = block
            text = para.text.strip()
            if not text:
                continue

            # 找 heading
            p_style = para.pPr.style if para.pPr and para.pPr.style else None
            # 修复：确保 style_name 是字符串而不是对象
            style_name = p_style.val if p_style is not None and hasattr(p_style, 'val') else (p_style if isinstance(p_style, str) else None)
            level = _parse_heading_level(style_name)

            if level is None:
                stack[-1].contents.append(text)
            else:
                # 解析节号和标题
                section_number, section_title = _parse_section_number(text)

                # 创建新的 section node
                while stack and stack[-1].level >= level:
                    stack.pop()
                node = _SectionNode(
                    title=text,
                    level=level,
                    section_number=section_number,
                    section_title=section_title
                )
                stack[-1].children.append(node)
                stack.append(node)

        elif block.tag.endswith('tbl'):  # 表格
            table = block
            # 用 python-docx 重新拿表格对象
            tbl = None
            for t in word_doc.tables:
                if t._tbl == table:
                    tbl = t
                    break
            if tbl:
                table_text = extract_table(tbl)
                if table_text:
                    stack[-1].contents.append(table_text)

    # --- flatten 成 Document 列表 ---
    docs: List[Document] = []
    global_chunk_index = [0]  # 使用列表以便在嵌套函数中修改

    def _flatten(node: _SectionNode, breadcrumb: List[str], breadcrumb_titles: List[str], hierarchy_stack: List[str]):
        """
        递归flatten节点树，生成带有丰富metadata的Documents

        Args:
            node: 当前节点
            breadcrumb: 面包屑路径（完整标题，如 ["3.2.1 LBS系统", "3.2.1.1 主图模式"]）
            breadcrumb_titles: 纯标题路径（如 ["LBS系统", "主图模式"]）
            hierarchy_stack: 节号栈（如 ["3.2.1", "3.2.1.1"]）
        """
        if node.contents:
            metadata = {'source': path.name}

            # 层级metadata
            if breadcrumb:
                metadata['section'] = ' > '.join(breadcrumb)  # 保留原有字段兼容性
                metadata['breadcrumb'] = ' > '.join(breadcrumb_titles)  # 纯标题路径
                metadata['section_level'] = len(breadcrumb)
                metadata['root_section'] = breadcrumb_titles[0] if breadcrumb_titles else None

                # 节号相关
                if node.section_number:
                    metadata['section_number'] = node.section_number
                if node.section_title:
                    metadata['section_title'] = node.section_title

                # 父节点信息（去掉最后一级，保留父路径）
                if len(breadcrumb) > 1:
                    metadata['parent_section'] = ' > '.join(breadcrumb_titles[:-1])
                else:
                    metadata['parent_section'] = None

            # 全局chunk索引
            metadata['global_chunk_index'] = global_chunk_index[0]
            global_chunk_index[0] += 1

            docs.append(
                Document(
                    page_content='\n'.join(node.contents),
                    metadata=metadata,
                )
            )

        for child in node.children:
            # 构建子节点的路径信息
            child_breadcrumb = breadcrumb + [child.title]
            child_breadcrumb_titles = breadcrumb_titles + [child.section_title or child.title]
            child_hierarchy_stack = hierarchy_stack + ([child.section_number] if child.section_number else [])

            _flatten(child, child_breadcrumb, child_breadcrumb_titles, child_hierarchy_stack)

    # Preface
    if root.contents:
        docs.append(
            Document(
                page_content='\n'.join(root.contents),
                metadata={
                    'source': path.name,
                    'section': 'preface',
                    'breadcrumb': '前言',
                    'section_level': 0,
                    'section_title': '前言',
                    'global_chunk_index': 0,
                },
            )
        )
        global_chunk_index[0] += 1

    for child in root.children:
        _flatten(child, [child.title], [child.section_title or child.title], [child.section_number] if child.section_number else [])

    return docs


def _load_fallback_docx(path: Path) -> List[Document]:
    loader = Docx2txtLoader(str(path))
    docs = loader.load()
    for d in docs:
        d.metadata.setdefault('source', path.name)
    return docs


def _analyze_font_sizes(pdf_path: Path) -> dict:
    """
    扫描 PDF 全文，分析字体大小分布，用于判断标题层级

    Returns:
        {
            'body_size': 正文字号（最常见的字号）,
            'heading_sizes': [一级标题字号, 二级标题字号, 三级标题字号],
            'size_to_level': {字号 -> 层级} 的映射
        }
    """
    from collections import Counter

    if not PDFPLUMBER_AVAILABLE:
        return {'body_size': 10, 'heading_sizes': [16, 14, 12], 'size_to_level': {}}

    font_sizes = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            chars = page.chars
            for char in chars:
                if 'size' in char and char['size']:
                    font_sizes.append(round(char['size'], 1))

    if not font_sizes:
        # 如果提取不到字体信息，使用默认值
        return {'body_size': 10, 'heading_sizes': [16, 14, 12], 'size_to_level': {}}

    # 统计字号分布
    size_counter = Counter(font_sizes)
    # 正文字号：最常见的字号
    body_size = size_counter.most_common(1)[0][0]

    # 找出所有大于正文的字号（可能是标题）
    unique_sizes = sorted(set(font_sizes), reverse=True)  # 从大到小排序
    heading_sizes = [s for s in unique_sizes if s > body_size]

    # 只取前3个层级（一级、二级、三级标题）
    heading_sizes = heading_sizes[:3]

    # 构建字号到层级的映射
    size_to_level = {}
    for i, size in enumerate(heading_sizes, start=1):
        size_to_level[size] = i

    return {
        'body_size': body_size,
        'heading_sizes': heading_sizes,
        'size_to_level': size_to_level,
    }


def _is_likely_heading(
    line_text: str,
    avg_font_size: float,
    body_size: float,
    size_to_level: dict,
    line_chars: List[dict] = None,
) -> Optional[int]:
    """
    判断一行文本是否可能是标题，并返回层级

    多重启发式规则：
    1. 字号大于正文 → 按字号映射层级
    2. 有数字编号（1. 1.1 等）→ 按层级判断
    3. 加粗字体 + 行短 → 可能是标题
    4. 特定关键词（第X章、附录等）→ 一级标题

    Args:
        line_text: 文本内容
        avg_font_size: 该行的平均字号
        body_size: 正文字号
        size_to_level: 字号到层级的映射
        line_chars: 该行的字符对象列表（用于检测加粗）

    Returns:
        如果是标题，返回层级 (1, 2, 3)；否则返回 None
    """
    text = line_text.strip()
    if not text:
        return None

    # 标题长度限制（太长的不太可能是标题）
    if len(text) > 200:
        return None

    # 规则1: 字号大于正文
    if avg_font_size > body_size:
        rounded_size = round(avg_font_size, 1)
        level = size_to_level.get(rounded_size)
        if level:
            return level

    # 规则2: 数字编号识别
    import re
    # 匹配 "1. ", "1.1 ", "1.1.1 ", "(1)", "第1章" 等
    number_patterns = [
        (r'^(\d+)\.\s+', lambda m: 1),  # "1. 标题" → 一级
        (r'^(\d+)\.(\d+)\s+', lambda m: 2),  # "1.1 标题" → 二级
        (r'^(\d+)\.(\d+)\.(\d+)\s+', lambda m: 3),  # "1.1.1 标题" → 三级
        (r'^第[一二三四五六七八九十\d]+章', lambda m: 1),  # "第一章" → 一级
        (r'^第[一二三四五六七八九十\d]+节', lambda m: 2),  # "第一节" → 二级
        (r'^\([一二三四五1-9]\)', lambda m: 2),  # "(一)" → 二级
    ]

    for pattern, level_func in number_patterns:
        match = re.match(pattern, text)
        if match:
            level = level_func(match)
            # 额外检查：如果文本太长，可能不是标题
            if len(text) < 100:
                return level

    # 规则3: 加粗 + 短文本
    if line_chars:
        # 检查是否大部分字符是加粗字体
        bold_chars = sum(1 for c in line_chars if 'Bold' in c.get('fontname', ''))
        if bold_chars > len(line_chars) * 0.5:  # 超过50%是加粗
            if len(text) < 80 and not text.endswith(('。', '，', '：', ';', ':')):
                return 2  # 默认二级标题

    # 规则4: 特定关键词
    heading_keywords = ['摘要', '概述', '前言', '引言', '背景', '功能说明', '接口定义', '附录']
    if text in heading_keywords or any(text.startswith(kw) for kw in heading_keywords):
        return 1

    return None


def _extract_table_from_pdfplumber(table_data) -> Optional[str]:
    """
    将 pdfplumber 提取的表格转换为结构化文本

    Args:
        table_data: pdfplumber 的 table 对象（二维列表）

    Returns:
        格式化的表格文本，用 " | " 分隔
    """
    if not table_data:
        return None

    rows = []
    for row in table_data:
        # 清理空值，转换为字符串
        cells = [str(cell).strip() if cell else '' for cell in row]
        # 跳过全空行
        if any(cells):
            rows.append(' | '.join(cells))

    return '\n'.join(rows) if rows else None


def _load_pdf_with_headings(path: Path) -> List[Document]:
    """
    使用 pdfplumber 解析 PDF，提取层级结构和表格

    功能：
    1. 通过字体大小识别标题层级
    2. 提取表格并保留结构
    3. 构建类似 DOCX 的 section 树状结构
    4. 返回带有 section metadata 的 Document 列表
    """
    if not PDFPLUMBER_AVAILABLE:
        raise ImportError("pdfplumber 未安装，无法使用高级 PDF 解析")

    # 第一步：分析字体分布
    font_analysis = _analyze_font_sizes(path)
    body_size = font_analysis['body_size']
    size_to_level = font_analysis['size_to_level']

    print(f"  PDF字体分析: 正文={body_size}pt, 标题层级={font_analysis['heading_sizes']}")

    # 第二步：构建 section 树
    root = _SectionNode(title='root', level=0)
    stack = [root]

    with pdfplumber.open(path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            # 提取文本（按行）
            text = page.extract_text()
            if text:
                lines = text.split('\n')

                for line in lines:
                    line = line.strip()
                    if not line:
                        continue

                    # 查找该行对应的字符（用于字号和加粗检测）
                    # 策略：匹配行首20个字符
                    line_prefix = line[:min(20, len(line))]
                    line_chars = []
                    for c in page.chars:
                        c_text = c.get('text', '')
                        if c_text and c_text in line_prefix:
                            line_chars.append(c)

                    # 计算平均字号
                    if line_chars:
                        avg_size = sum(c.get('size', body_size) for c in line_chars) / len(line_chars)
                    else:
                        avg_size = body_size

                    # 判断是否是标题（传入字符信息用于加粗检测）
                    level = _is_likely_heading(
                        line, avg_size, body_size, size_to_level, line_chars
                    )

                    if level:
                        # 是标题，解析节号和标题
                        section_number, section_title = _parse_section_number(line)

                        # 创建新的 section node
                        while stack and stack[-1].level >= level:
                            stack.pop()
                        node = _SectionNode(
                            title=line,
                            level=level,
                            section_number=section_number,
                            section_title=section_title
                        )
                        stack[-1].children.append(node)
                        stack.append(node)
                    else:
                        # 是正文，添加到当前 section
                        stack[-1].contents.append(line)

            # 提取表格
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    table_text = _extract_table_from_pdfplumber(table)
                    if table_text:
                        stack[-1].contents.append(table_text)

    # 第三步：flatten 成 Document 列表
    docs: List[Document] = []
    global_chunk_index = [0]  # 使用列表以便在嵌套函数中修改

    def _flatten(node: _SectionNode, breadcrumb: List[str], breadcrumb_titles: List[str], hierarchy_stack: List[str]):
        """
        递归flatten节点树，生成带有丰富metadata的Documents

        Args:
            node: 当前节点
            breadcrumb: 面包屑路径（完整标题）
            breadcrumb_titles: 纯标题路径
            hierarchy_stack: 节号栈
        """
        if node.contents:
            metadata = {'source': path.name}

            # 层级metadata
            if breadcrumb:
                metadata['section'] = ' > '.join(breadcrumb)  # 保留原有字段兼容性
                metadata['breadcrumb'] = ' > '.join(breadcrumb_titles)  # 纯标题路径
                metadata['section_level'] = len(breadcrumb)
                metadata['root_section'] = breadcrumb_titles[0] if breadcrumb_titles else None

                # 节号相关
                if node.section_number:
                    metadata['section_number'] = node.section_number
                if node.section_title:
                    metadata['section_title'] = node.section_title

                # 父节点信息（去掉最后一级，保留父路径）
                if len(breadcrumb) > 1:
                    metadata['parent_section'] = ' > '.join(breadcrumb_titles[:-1])
                else:
                    metadata['parent_section'] = None

            # 全局chunk索引
            metadata['global_chunk_index'] = global_chunk_index[0]
            global_chunk_index[0] += 1

            docs.append(
                Document(
                    page_content='\n'.join(node.contents),
                    metadata=metadata,
                )
            )

        for child in node.children:
            child_breadcrumb = breadcrumb + [child.title]
            child_breadcrumb_titles = breadcrumb_titles + [child.section_title or child.title]
            child_hierarchy_stack = hierarchy_stack + ([child.section_number] if child.section_number else [])

            _flatten(child, child_breadcrumb, child_breadcrumb_titles, child_hierarchy_stack)

    # Preface（无标题的前言部分）
    if root.contents:
        docs.append(
            Document(
                page_content='\n'.join(root.contents),
                metadata={
                    'source': path.name,
                    'section': 'preface',
                    'breadcrumb': '前言',
                    'section_level': 0,
                    'section_title': '前言',
                    'global_chunk_index': 0,
                },
            )
        )
        global_chunk_index[0] += 1

    for child in root.children:
        _flatten(child, [child.title], [child.section_title or child.title], [child.section_number] if child.section_number else [])

    return docs


def _load_excel_as_rows(path: Path) -> List[Document]:
    """
    加载Excel文件，按行切分成chunks

    对于有层级列（level1, level2...）的表格，会forward-fill空白单元格
    把每一行转成可读的文本格式

    支持多级表头的Excel：会智能合并主表头和子表头
    """
    if not PANDAS_AVAILABLE:
        raise ImportError("需要安装 pandas 和 openpyxl 来加载 Excel 文件")

    # 读取Excel，尝试检测多级表头
    df = pd.read_excel(path, engine='openpyxl')

    # 检测并处理多级表头
    # 如果第一行数据看起来像表头（大部分是文字描述），则可能是多级表头
    if len(df) > 0:
        first_row = df.iloc[0]
        # 如果第一行有很多文字内容（超过50%的单元格有长文本），可能是子表头
        text_cells = sum(1 for val in first_row if pd.notna(val) and isinstance(val, str) and len(str(val)) > 10)

        if text_cells > len(df.columns) * 0.3:  # 如果超过30%是长文本，可能是子表头
            # 重新读取，使用前2行作为表头
            df_multi = pd.read_excel(path, engine='openpyxl', header=[0, 1])

            # 合并多级列名
            new_columns = []
            for col in df_multi.columns:
                if isinstance(col, tuple):
                    # 合并非空的列名部分
                    parts = [str(c).strip() for c in col if not str(c).startswith('Unnamed')]
                    if parts:
                        new_columns.append(' - '.join(parts))
                    else:
                        new_columns.append(f"列{len(new_columns)}")
                else:
                    new_columns.append(str(col))

            df_multi.columns = new_columns
            df = df_multi

    # 如果还有Unnamed列，尝试智能重命名
    new_columns = []
    last_named_col = "未分类"

    for col in df.columns:
        col_str = str(col)
        if col_str.startswith('Unnamed:'):
            # Unnamed列用前一个有名字的列作为前缀
            new_columns.append(f"{last_named_col}-子列{col_str.split(':')[1]}")
        else:
            new_columns.append(col_str)
            last_named_col = col_str

    df.columns = new_columns

    # 检测并处理层级列（以level开头的列）
    level_cols = [col for col in df.columns if str(col).lower().startswith('level')]
    if level_cols:
        # Forward fill层级列的空值（处理合并单元格）
        df[level_cols] = df[level_cols].ffill()

    docs: List[Document] = []

    # 遍历每一行
    for idx, row in df.iterrows():
        # 构建文本内容
        parts = []

        # 提取层级值（用于metadata）
        hierarchy_values = {}
        hierarchy_path_parts = []

        # 1. 如果有层级列，先构建层级路径
        if level_cols:
            hierarchy = []
            for level_idx, col in enumerate(level_cols, start=1):
                val = row[col]
                if pd.notna(val) and str(val).strip() and str(val).strip() != '/':
                    val_str = str(val).strip()
                    hierarchy.append(val_str)
                    # 保存到独立字段 level1, level2, level3, level4
                    hierarchy_values[f'level{level_idx}'] = val_str
                    hierarchy_path_parts.append(val_str)
                else:
                    hierarchy_values[f'level{level_idx}'] = None

            if hierarchy:
                parts.append(f"层级: {' > '.join(hierarchy)}")

        # 2. 添加其他列的内容
        for col in df.columns:
            if col not in level_cols:
                val = row[col]
                if pd.notna(val) and str(val).strip():
                    parts.append(f"{col}: {str(val).strip()}")

        # 如果这行有实质内容，创建Document
        if parts:
            content = '\n'.join(parts)
            metadata = {
                'source': path.name,
                'row_number': int(idx) + 2,  # +2 因为Excel行号从1开始，且有表头
                'file_type': path.suffix.lower(),
                'global_chunk_index': int(idx),
            }

            # 添加层级metadata
            if hierarchy_values:
                # 添加各个level字段
                metadata.update(hierarchy_values)

                # 构建hierarchy_path
                if hierarchy_path_parts:
                    metadata['hierarchy_path'] = ' > '.join(hierarchy_path_parts)
                    metadata['breadcrumb'] = ' > '.join(hierarchy_path_parts)
                    metadata['section_level'] = len(hierarchy_path_parts)
                    metadata['root_section'] = hierarchy_path_parts[0] if hierarchy_path_parts else None

                    # 父节点（去掉最后一级）
                    if len(hierarchy_path_parts) > 1:
                        metadata['parent_section'] = ' > '.join(hierarchy_path_parts[:-1])
                    else:
                        metadata['parent_section'] = None

            docs.append(Document(page_content=content, metadata=metadata))

    return docs


def _load_fallback_pdf(path: Path) -> List[Document]:
    """降级 PDF 加载器（按页提取）"""
    loader = PyPDFLoader(str(path))
    docs = loader.load()
    for d in docs:
        d.metadata.setdefault('source', path.name)
    return docs


def _load_with_metadata(path: Path) -> List[Document]:
    """
    加载文档并添加元数据（doc_type, file_type等）

    策略：
    - DOCX: 尝试 _load_docx_with_headings，失败则降级到 Docx2txtLoader
    - PDF: 尝试 _load_pdf_with_headings (需要 pdfplumber)，失败则降级到 PyPDFLoader
    - Excel: 使用 _load_excel_as_rows 按行加载（需要 pandas）
    """
    suffix = path.suffix.lower()

    # 加载文档
    if suffix == '.docx':
        try:
            docs = _load_docx_with_headings(path)
        except Exception as e:
            print(f"  ⚠️  DOCX高级解析失败，使用降级加载器: {e}")
            docs = _load_fallback_docx(path)

    elif suffix == '.pdf':
        if PDFPLUMBER_AVAILABLE:
            try:
                docs = _load_pdf_with_headings(path)
            except Exception as e:
                print(f"  ⚠️  PDF高级解析失败，使用降级加载器: {e}")
                docs = _load_fallback_pdf(path)
        else:
            docs = _load_fallback_pdf(path)

    elif suffix in ['.xlsx', '.xls']:
        if PANDAS_AVAILABLE:
            try:
                docs = _load_excel_as_rows(path)
            except Exception as e:
                print(f"  ⚠️  Excel加载失败: {e}")
                raise ValueError(f"无法加载Excel文件 {path.name}")
        else:
            print(f"  ⚠️  跳过Excel文件（pandas未安装）: {path.name}")
            raise ValueError(f"pandas未安装，无法加载 {path.name}")

    else:
        # 其他格式，使用默认加载器
        loader = _build_loader(path)
        docs = loader.load()
        for d in docs:
            d.metadata.setdefault('source', path.name)

    # 检测文档类型并添加元数据
    doc_type = detect_doc_type(path)
    for d in docs:
        d.metadata['doc_type'] = doc_type
        d.metadata['file_type'] = suffix  # '.pdf' 或 '.docx'
        d.metadata['doc_path'] = str(path)

    return docs


def load_documents(docs_dir: str = 'data/documents') -> List[Document]:
    base_path = Path(docs_dir)
    if not base_path.exists():
        raise FileNotFoundError(f'{docs_dir} does not exist.')

    all_docs: List[Document] = []
    for path in base_path.rglob('*'):
        if path.is_dir():
            continue
        try:
            docs = _load_with_metadata(path)
        except ValueError:
            continue
        all_docs.extend(docs)

    print(f'Loaded document chunks: {len(all_docs)}')
    return all_docs


def load_single_document(file_path: str) -> List[Document]:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f'{file_path} does not exist.')

    docs = _load_with_metadata(path)
    print(f'{path.name} chunks: {len(docs)}')
    return docs
